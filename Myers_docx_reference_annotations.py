"""DOCX rendering helpers for a reference-based structured diff.

This writer keeps the reference text as the base document and wraps only the
changed spans with compact inline annotations:
- {+ text +} for insertions
- {- text -} for deletions
- {~ old -> new ~} for replacements

The goal is a small and readable document that is easy to backtrack.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from Myers_diff import DiffChange, DiffOutput, ReplaceTokenDiff, compute_diff, diff_replace_block

__all__ = [
    "build_reference_annotated_document",
    "save_reference_annotated_docx",
    "save_reference_annotated_docx_from_text",
]


# One place to tune the final DOCX styling.
REFERENCE_ANNOTATION_THEME = {
    "font_name": "Calibri",
    "font_size_pt": 10,
    "line_number_color": "808080",
    "line_number_bold": True,
    "equal": {
        "font_color": "000000",
        "fill_color": None,
        "bold": False,
    },
    "delete": {
        "font_color": "C00000",
        "fill_color": "FDE9E7",
        "bold": True,
    },
    "insert": {
        "font_color": "008000",
        "fill_color": "E6F4EA",
        "bold": True,
    },
    "marker": {
        "font_color": "808080",
        "bold": True,
    },
}


# Keep the page compact so the annotations stay readable.
REFERENCE_PAGE_MARGINS = {
    "top": Cm(0.5),
    "bottom": Cm(0.5),
    "left": Cm(0.5),
    "right": Cm(0.5),
}

BLACK = RGBColor.from_string("000000")
GREEN = RGBColor.from_string("008000")
RED = RGBColor.from_string("C00000")
GREY = RGBColor.from_string("808080")

# Keep same-type token runs bounded so long boilerplate does not collapse into
# one giant annotation span.
MAX_COALESCED_SEGMENT_CHARS = 96
SEMANTIC_BREAK_TOKENS = ("|", ":", ";", "/", "(", ")", "[", "]", "{", "}")


def _rgb(hex_value: str) -> RGBColor:
    """Convert a hex string like 'C00000' into a Word RGB color."""

    return RGBColor.from_string(hex_value)


def _set_run_fill(run, fill_hex: str | None) -> None:
    """Apply a custom background fill to a Word run."""

    if not fill_hex:
        return

    run_properties = run._element.get_or_add_rPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill_hex)
    run_properties.append(shading)


def _style_run(
    run,
    *,
    color: RGBColor | None = None,
    bold: bool = False,
    fill_hex: str | None = None,
) -> None:
    """Apply the shared visual style to one run."""

    run.font.name = REFERENCE_ANNOTATION_THEME["font_name"]
    run.font.size = Pt(REFERENCE_ANNOTATION_THEME["font_size_pt"])
    if color is not None:
        run.font.color.rgb = color
    run.font.bold = bold
    _set_run_fill(run, fill_hex)


def _add_run(
    paragraph,
    text: str,
    *,
    color: RGBColor | None = None,
    bold: bool = False,
    fill_hex: str | None = None,
):
    """Add one styled run to a paragraph and return it."""

    run = paragraph.add_run(text)
    _style_run(run, color=color, bold=bold, fill_hex=fill_hex)
    return run


def _prepare_paragraph(paragraph) -> None:
    """Keep paragraphs compact so the structured annotations stay tight."""

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def _add_line_number(paragraph, number: int | None) -> None:
    """Write an aligned line number at the start of a paragraph."""

    if number is None:
        return

    _add_run(
        paragraph,
        f"{number:>4} ",
        color=_rgb(REFERENCE_ANNOTATION_THEME["line_number_color"]),
        bold=REFERENCE_ANNOTATION_THEME["line_number_bold"],
    )


def _add_marker_open(paragraph, marker: str) -> None:
    """Write the left side of a structured annotation."""

    _add_run(
        paragraph,
        marker,
        color=_rgb(REFERENCE_ANNOTATION_THEME["marker"]["font_color"]),
        bold=REFERENCE_ANNOTATION_THEME["marker"]["bold"],
    )


def _add_marker_close(paragraph, marker: str) -> None:
    """Write the right side of a structured annotation."""

    _add_run(
        paragraph,
        marker,
        color=_rgb(REFERENCE_ANNOTATION_THEME["marker"]["font_color"]),
        bold=REFERENCE_ANNOTATION_THEME["marker"]["bold"],
    )


def _write_structured_annotation(
    paragraph,
    kind: str,
    text: str,
    *,
    replacement_text: str | None = None,
) -> None:
    """Wrap one changed span in a compact structured annotation.

    The marker tells the reader what changed, while the colored text keeps the
    document small and easy to scan.
    """

    if kind == "delete":
        _add_marker_open(paragraph, "{- ")
        _add_run(
            paragraph,
            text,
            color=RED,
            bold=REFERENCE_ANNOTATION_THEME["delete"]["bold"],
            fill_hex=REFERENCE_ANNOTATION_THEME["delete"]["fill_color"],
        )
        _add_marker_close(paragraph, " -}")
        return

    if kind == "insert":
        _add_marker_open(paragraph, "{+ ")
        _add_run(
            paragraph,
            text,
            color=GREEN,
            bold=REFERENCE_ANNOTATION_THEME["insert"]["bold"],
            fill_hex=REFERENCE_ANNOTATION_THEME["insert"]["fill_color"],
        )
        _add_marker_close(paragraph, " +}")
        return

    if kind == "replace":
        _add_marker_open(paragraph, "{~ ")
        _add_run(
            paragraph,
            text,
            color=RED,
            bold=REFERENCE_ANNOTATION_THEME["delete"]["bold"],
            fill_hex=REFERENCE_ANNOTATION_THEME["delete"]["fill_color"],
        )
        _add_run(paragraph, " -> ", color=_rgb(REFERENCE_ANNOTATION_THEME["marker"]["font_color"]), bold=True)
        _add_run(
            paragraph,
            replacement_text or "",
            color=GREEN,
            bold=REFERENCE_ANNOTATION_THEME["insert"]["bold"],
            fill_hex=REFERENCE_ANNOTATION_THEME["insert"]["fill_color"],
        )
        _add_marker_close(paragraph, " ~}")
        return

    raise ValueError(f"Unsupported structured kind: {kind}")


def _write_plain_line(document: Any, number: int | None, text: str) -> None:
    """Write an unchanged reference line."""

    paragraph = document.add_paragraph()
    _prepare_paragraph(paragraph)
    _add_line_number(paragraph, number)
    if text:
        _add_run(paragraph, text, color=BLACK, bold=False)


def _write_delete_line(document: Any, number: int | None, text: str) -> None:
    """Write a deleted reference line with a deletion annotation."""

    paragraph = document.add_paragraph()
    _prepare_paragraph(paragraph)
    _add_line_number(paragraph, number)
    _write_structured_annotation(paragraph, "delete", text)


def _write_insert_line(document: Any, number: int | None, text: str) -> None:
    """Write an inserted line as a compact structured annotation."""

    paragraph = document.add_paragraph()
    _prepare_paragraph(paragraph)
    _add_line_number(paragraph, number)
    _write_structured_annotation(paragraph, "insert", text)


def _append_token_text(paragraph, text: str, *, color: RGBColor | None = None, fill_hex: str | None = None) -> None:
    """Append raw token text without wrapping it in a structured marker."""

    if text:
        _add_run(paragraph, text, color=color, fill_hex=fill_hex)


def _has_semantic_break(text: str) -> bool:
    """Return True when a token segment looks like a field boundary."""

    return any(marker in text for marker in SEMANTIC_BREAK_TOKENS)


def _coalesce_token_edits(edits) -> list[tuple[str, str]]:
    """Merge adjacent token edits that share the same type.

    This keeps whitespace and punctuation attached to the nearby change so the
    rendered annotation stays compact and readable.
    """

    segments: list[tuple[str, str]] = []

    for edit in edits:
        if edit.type == "equal":
            segment_type = "equal"
            segment_text = edit.old_line.text if edit.old_line is not None else ""
        elif edit.type == "delete":
            segment_type = "delete"
            segment_text = edit.old_line.text if edit.old_line is not None else ""
        elif edit.type == "insert":
            segment_type = "insert"
            segment_text = edit.new_line.text if edit.new_line is not None else ""
        else:
            raise ValueError(f"Unsupported edit type in token diff: {edit.type}")

        if segments and segments[-1][0] == segment_type:
            previous_type, previous_text = segments[-1]

            # Keep same-type change chunks short and stop at field-like anchors.
            should_merge = True
            if segment_type in {"delete", "insert"}:
                should_merge = (
                    len(previous_text) + len(segment_text) <= MAX_COALESCED_SEGMENT_CHARS
                    and not _has_semantic_break(previous_text)
                    and not _has_semantic_break(segment_text)
                )

            if should_merge:
                segments[-1] = (previous_type, previous_text + segment_text)
            else:
                segments.append((segment_type, segment_text))
        else:
            segments.append((segment_type, segment_text))

    return segments


def _write_token_level_replace_line(paragraph, token_diff: ReplaceTokenDiff) -> None:
    """Render a single-line replacement using the token-level diff.

    Equal tokens stay plain. Delete/insert pairs are merged into a single
    {~ old -> new ~} annotation so the substitution is obvious.
    """

    segments = _coalesce_token_edits(token_diff.token_diff.edits)
    index = 0

    while index < len(segments):
        segment_type, segment_text = segments[index]

        if segment_type == "equal":
            _append_token_text(paragraph, segment_text, color=BLACK)
            index += 1
            continue

        if segment_type == "delete" and index + 1 < len(segments) and segments[index + 1][0] == "insert":
            new_text = segments[index + 1][1]
            old_text = segment_text
            _write_structured_annotation(paragraph, "replace", old_text, replacement_text=new_text)
            index += 2
            continue

        if segment_type == "insert" and index + 1 < len(segments) and segments[index + 1][0] == "delete":
            old_text = segments[index + 1][1]
            new_text = segment_text
            _write_structured_annotation(paragraph, "replace", old_text, replacement_text=new_text)
            index += 2
            continue

        if segment_type == "delete":
            old_text = segment_text
            _write_structured_annotation(paragraph, "delete", old_text)
            index += 1
            continue

        if segment_type == "insert":
            new_text = segment_text
            _write_structured_annotation(paragraph, "insert", new_text)
            index += 1
            continue

        raise ValueError(f"Unsupported edit type in token diff: {segment_type}")


def _write_token_level_replace_pair(document: Any, old_line, new_line, *, line_number: int | None) -> None:
    """Render one replace pair as a small semantic chunk.

    This keeps the reference text visible as plain context and only wraps the
    actual delta, which makes long document-style lines much easier to review.
    """

    paragraph = document.add_paragraph()
    _prepare_paragraph(paragraph)
    _add_line_number(paragraph, line_number)

    pair_block = DiffChange(type="replace", old_lines=[old_line], new_lines=[new_line])
    pair_token_diff = ReplaceTokenDiff(block=pair_block, token_diff=diff_replace_block(pair_block))
    _write_token_level_replace_line(paragraph, pair_token_diff)


def _write_replacement_block(document: Any, block: DiffChange) -> None:
    """Render a replace block using the smallest readable representation.

    For single-line replace blocks, the token diff is rendered inline.
    For multi-line replace blocks, each old/new pair becomes its own chunk so
    the reader can focus on one local substitution at a time.
    """

    if len(block.old_lines) == 1 and len(block.new_lines) == 1:
        _write_token_level_replace_pair(
            document,
            block.old_lines[0],
            block.new_lines[0],
            line_number=block.old_lines[0].number,
        )
        return

    pair_count = max(len(block.old_lines), len(block.new_lines))

    for index in range(pair_count):
        old_line = block.old_lines[index] if index < len(block.old_lines) else None
        new_line = block.new_lines[index] if index < len(block.new_lines) else None

        line_number = old_line.number if old_line is not None else (new_line.number if new_line is not None else None)

        if old_line is not None and new_line is not None:
            _write_token_level_replace_pair(
                document,
                old_line,
                new_line,
                line_number=line_number,
            )
        elif old_line is not None:
            paragraph = document.add_paragraph()
            _prepare_paragraph(paragraph)
            _add_line_number(paragraph, line_number)
            _write_structured_annotation(paragraph, "delete", old_line.text)
        elif new_line is not None:
            paragraph = document.add_paragraph()
            _prepare_paragraph(paragraph)
            _add_line_number(paragraph, line_number)
            _write_structured_annotation(paragraph, "insert", new_line.text)


def _write_header(document: Any, result: DiffOutput, title: str) -> None:
    """Add a compact header and legend for the structured diff."""

    document.add_heading(title, level=1)

    summary = document.add_paragraph()
    _prepare_paragraph(summary)
    _add_run(summary, f"Edit distance: {result.edit_distance}", bold=True)
    _add_run(summary, f" | blocks: {len(result.change_blocks)}", color=GREY)

    legend = document.add_paragraph()
    _prepare_paragraph(legend)
    _add_run(legend, "Legend: ", color=GREY, bold=True)
    _add_run(legend, "{+ insert +} ", color=GREEN, bold=True)
    _add_run(legend, "{- delete -} ", color=RED, bold=True)
    _add_run(legend, "{~ replace old -> new ~}", color=GREY, bold=True)


def build_reference_annotated_document(result: DiffOutput, *, title: str = "Myers Reference Annotated Diff") -> Any:
    """Build a Word document that keeps the reference text as the base."""

    document = Document()

    for section in document.sections:
        section.top_margin = REFERENCE_PAGE_MARGINS["top"]
        section.bottom_margin = REFERENCE_PAGE_MARGINS["bottom"]
        section.left_margin = REFERENCE_PAGE_MARGINS["left"]
        section.right_margin = REFERENCE_PAGE_MARGINS["right"]

    normal_style: Any = document.styles["Normal"]
    normal_style.font.name = REFERENCE_ANNOTATION_THEME["font_name"]
    normal_style.font.size = Pt(REFERENCE_ANNOTATION_THEME["font_size_pt"])

    _write_header(document, result, title)

    for block in result.change_blocks:
        if block.type == "equal":
            for line in block.old_lines:
                _write_plain_line(document, line.number, line.text)
        elif block.type == "delete":
            for line in block.old_lines:
                _write_delete_line(document, line.number, line.text)
        elif block.type == "insert":
            for line in block.new_lines:
                _write_insert_line(document, line.number, line.text)
        elif block.type == "replace":
            _write_replacement_block(document, block)
        else:
            raise ValueError(f"Unsupported block type: {block.type}")

        # Keep a small gap between blocks so the annotations stay easy to scan.
        document.add_paragraph()

    return document


def save_reference_annotated_docx(
    result: DiffOutput,
    output_path: str | Path,
    *,
    title: str = "Myers Reference Annotated Diff",
) -> Path:
    """Save the structured reference-annotated diff as a .docx file."""

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document = build_reference_annotated_document(result, title=title)
    # ensure the text is using the right front color, type and size before saving
    document.styles['Normal'].font.name = REFERENCE_ANNOTATION_THEME["font_name"]
    document.styles['Normal'].font.size = Pt(REFERENCE_ANNOTATION_THEME["font_size_pt"])
    document.save(str(path))
    return path


def save_reference_annotated_docx_from_text(
    reference_text: str,
    modified_text: str,
    output_path: str | Path,
    *,
    title: str = "Myers Reference Annotated Diff",
) -> Path:
    """Compute the diff first, then save the structured DOCX output."""

    result = compute_diff(reference_text, modified_text)
    return save_reference_annotated_docx(result, output_path, title=title)