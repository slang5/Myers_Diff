"""DOCX rendering helpers for the Myers diff output.

This module keeps the presentation layer separate from the Myers algorithm:
- compute the line-level diff in Myers_diff.py
- group replace blocks there
- render the final output here as a stacked, GitHub-like diff view
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from Myers_diff import DiffChange, DiffOutput, ReplaceTokenDiff, build_replace_token_diffs, compute_diff

__all__ = [
    "build_diff_document",
    "save_diff_docx",
    "save_diff_docx_from_text",
]

# One place to tune the final DOCX styling.
# Adjust these hex values to change the visual theme quickly.
DIFF_RENDER_THEME = {
    "font_name": "Calibri",
    "font_size_pt": 10,
    "line_number_color": "808080",
    "line_number_bold": True,
    "equal": {
        "font_color": "000000",
        "fill_color": None,
        "bold": False,
        "prefix": "",
    },
    "delete": {
        "font_color": "C00000",
        "fill_color": "FDE9E7",
        "bold": True,
        "prefix": "- ",
    },
    "insert": {
        "font_color": "008000",
        "fill_color": "E6F4EA",
        "bold": True,
        "prefix": "+ ",
    },
    "replace": {
        "font_color": "B59A00",
        "fill_color": "FFF8CC",
        "bold": True,
        "prefix": "~ ",
    },
}

# Define word page margins as constants for easy adjustment.
word_page_margins = {
    "top": Cm(0.5),
    "bottom": Cm(0.5),
    "left": Cm(0.5),
    "right": Cm(0.5),
}

GREEN = RGBColor.from_string("008000")
RED = RGBColor.from_string("C00000")
BLACK = RGBColor.from_string("000000")
GREY = RGBColor.from_string("808080")
TABLE_STYLE = "Table Grid"


def _rgb(hex_value: str) -> RGBColor:
    """Convert a hex string like 'C00000' into a Word RGB color."""

    return RGBColor.from_string(hex_value)


def _set_run_fill(run, fill_hex: str | None) -> None:
    """Apply a custom background fill to a Word run."""

    if not fill_hex:
        return

    run_properties = run._element.get_or_add_rPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill_hex)
    run_properties.append(shading)

def _iter_replace_token_diffs(result: DiffOutput) -> Iterable[ReplaceTokenDiff]:
    """Yield token-level diffs only for replace blocks."""

    for item in build_replace_token_diffs(result):
        yield item

def _style_run_mono(
    run,
    *,
    color: RGBColor | None = None,
    bold: bool = False,
    fill_hex: str | None = None,
    highlight: bool = False,
) -> None:
    """Apply a consistent monospace style to a Word run."""

    run.font.name = DIFF_RENDER_THEME["font_name"]
    run.font.size = Pt(DIFF_RENDER_THEME["font_size_pt"])
    if color is not None:
        run.font.color.rgb = color
    run.font.bold = bold
    if fill_hex is None and highlight:
        fill_hex = DIFF_RENDER_THEME["replace"]["fill_color"]
    _set_run_fill(run, fill_hex)


def _add_run_mono(
    paragraph,
    text: str,
    *,
    color: RGBColor | None = None,
    bold: bool = False,
    fill_hex: str | None = None,
    highlight: bool = False,
):
    """Add a styled run to a paragraph and return it."""

    run = paragraph.add_run(text)
    _style_run_mono(run, color=color, bold=bold, fill_hex=fill_hex, highlight=highlight)
    return run

def _add_vertical_diff_header(document: Any, result: DiffOutput, title: str) -> None:
    """Add a compact header for the vertical diff view."""

    document.add_heading(title, level=1)

    summary = document.add_paragraph()
    _add_run_mono(summary, f"Edit distance: {result.edit_distance}", bold=True)
    _add_run_mono(summary, f" | blocks: {len(result.change_blocks)}", color=GREY)

    note = document.add_paragraph()
    _add_run_mono(note, "Vertical diff view", color=GREY, bold=True)


def _write_vertical_line(
    document: Any,
    *,
    number: int | None,
    text: str,
    kind: str = "equal",
    prefix: str | None = None,
) -> None:
    """Write one diff line in a simple stacked format."""

    paragraph = document.add_paragraph()

    style = DIFF_RENDER_THEME[kind]
    font_color = _rgb(style["font_color"])
    fill_color = style["fill_color"]
    prefix_text = style["prefix"] if prefix is None else prefix

    if number is not None:
        _add_run_mono(
            paragraph,
            f"{number:>4} ",
            color=_rgb(DIFF_RENDER_THEME["line_number_color"]),
            bold=DIFF_RENDER_THEME["line_number_bold"],
        )

    if prefix_text:
        _add_run_mono(paragraph, prefix_text, color=font_color, bold=style["bold"], fill_hex=fill_color)

    if text:
        _add_run_mono(paragraph, text, color=font_color, bold=style["bold"], fill_hex=fill_color)


def _render_vertical_equal_block(document: Any, block: DiffChange) -> None:
    """Render unchanged lines in a simple stacked layout."""

    for line in block.old_lines:
        _write_vertical_line(document, number=line.number, text=line.text, kind="equal")


def _render_vertical_delete_block(document: Any, block: DiffChange) -> None:
    """Render deleted lines in red."""

    for line in block.old_lines:
        _write_vertical_line(document, number=line.number, text=line.text, kind="delete")


def _render_vertical_insert_block(document: Any, block: DiffChange) -> None:
    """Render inserted lines in green."""

    for line in block.new_lines:
        _write_vertical_line(document, number=line.number, text=line.text, kind="insert")


def _render_vertical_replace_block(document: Any, token_diff: ReplaceTokenDiff) -> None:
    """Render a replace block showing intra-line changes."""

    # We need to reconstruct the lines based on the token edits.
    # A single ReplaceTokenDiff covers the whole block of potentially multiple old and new lines.

    # 1. Gather all tokens for the "old" side (delete + equal)
    old_paragraphs = []
    current_old_p = document.add_paragraph()
    _add_run_mono(current_old_p, "- ", color=RED, bold=DIFF_RENDER_THEME["delete"]["bold"], fill_hex=DIFF_RENDER_THEME["delete"]["fill_color"])
    old_paragraphs.append(current_old_p)
    
    for edit in token_diff.token_diff.edits:
        if edit.type in ("equal", "delete") and edit.old_line is not None:
            text = edit.old_line.text
            # split text on newlines to handle multi-line token blocks
            parts = text.split("\n")
            for i, part in enumerate(parts):
                if part:
                    color = RED if edit.type == "delete" else _rgb(DIFF_RENDER_THEME["replace"]["font_color"])
                    fill = DIFF_RENDER_THEME["delete"]["fill_color"] if edit.type == "delete" else DIFF_RENDER_THEME["replace"]["fill_color"]
                    _add_run_mono(current_old_p, part, color=color, bold=True, fill_hex=fill)
                if i < len(parts) - 1:
                    current_old_p = document.add_paragraph()
                    _add_run_mono(current_old_p, "- ", color=RED, bold=DIFF_RENDER_THEME["delete"]["bold"], fill_hex=DIFF_RENDER_THEME["delete"]["fill_color"])
                    old_paragraphs.append(current_old_p)

    # 2. Gather all tokens for the "new" side (insert + equal)
    new_paragraphs = []
    current_new_p = document.add_paragraph()
    _add_run_mono(current_new_p, "+ ", color=GREEN, bold=DIFF_RENDER_THEME["insert"]["bold"], fill_hex=DIFF_RENDER_THEME["insert"]["fill_color"])
    new_paragraphs.append(current_new_p)

    for edit in token_diff.token_diff.edits:
        if edit.type in ("equal", "insert") and edit.new_line is not None:
            text = edit.new_line.text
            parts = text.split("\n")
            for i, part in enumerate(parts):
                if part:
                    color = GREEN if edit.type == "insert" else _rgb(DIFF_RENDER_THEME["replace"]["font_color"])
                    fill = DIFF_RENDER_THEME["insert"]["fill_color"] if edit.type == "insert" else DIFF_RENDER_THEME["replace"]["fill_color"]
                    _add_run_mono(current_new_p, part, color=color, bold=True, fill_hex=fill)
                if i < len(parts) - 1:
                    current_new_p = document.add_paragraph()
                    _add_run_mono(current_new_p, "+ ", color=GREEN, bold=DIFF_RENDER_THEME["insert"]["bold"], fill_hex=DIFF_RENDER_THEME["insert"]["fill_color"])
                    new_paragraphs.append(current_new_p)


def _build_vertical_diff(document: Any, result: DiffOutput) -> None:
    """Render the diff as stacked lines instead of a 4-column table."""

    replace_token_diffs = iter(_iter_replace_token_diffs(result))

    for block in result.change_blocks:
        if block.type == "equal":
            _render_vertical_equal_block(document, block)
        elif block.type == "delete":
            _render_vertical_delete_block(document, block)
        elif block.type == "insert":
            _render_vertical_insert_block(document, block)
        elif block.type == "replace":
            token_diff = next(replace_token_diffs)
            _render_vertical_replace_block(document, token_diff)
        else:
            raise ValueError(f"Unsupported block type: {block.type}")

        document.add_paragraph()


def build_diff_document(result: DiffOutput, *, title: str = "Myers Diff Output") -> Any:
    """Build a Word document that shows a stacked GitHub-like diff."""

    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = word_page_margins["top"]
        section.bottom_margin = word_page_margins["bottom"]
        section.left_margin = word_page_margins["left"]
        section.right_margin = word_page_margins["right"]

    _add_vertical_diff_header(document, result, title)
    _build_vertical_diff(document, result)

    return document


def save_diff_docx(result: DiffOutput, output_path: str | Path, *, title: str = "Myers Diff Output") -> Path:
    """Save an already computed diff result as a .docx file."""

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document = build_diff_document(result, title=title)
    document.save(str(path))
    return path


def save_diff_docx_from_text(
    old_text: str,
    new_text: str,
    output_path: str | Path,
    *,
    title: str = "Myers Diff Output",
) -> Path:
    """Compute the diff first, then save it as a .docx file."""

    result = compute_diff(old_text, new_text)
    return save_diff_docx(result, output_path, title=title)